from docx import Document

class ResumePostProcessor:
    @staticmethod
    def fix_page_breaks(docx_path: str) -> str:
        """
        Ensures headings and critical sections stay together on the same page.
        This modifies the file in place.
        """
        doc = Document(docx_path)
        
        for paragraph in doc.paragraphs:
            # Simple heuristic to keep headings together with the next paragraph
            if paragraph.style.name.startswith("Heading"):
                paragraph.paragraph_format.keep_with_next = True
                
        doc.save(docx_path)
        return docx_path
    
    @staticmethod
    def remove_empty_bullets(docx_path: str) -> str:
        """
        Removes bullet/list paragraphs that contain no text.
        Useful after docxtpl rendering when variables are empty.
        """
        doc = Document(docx_path)

        for paragraph in list(doc.paragraphs):
            text = paragraph.text.strip()

            style_name = paragraph.style.name.lower()

            # Detect bullet/list styles
            is_bullet = "list" in style_name or "bullet" in style_name

            if is_bullet and text == "":
                paragraph._element.getparent().remove(paragraph._element)

        doc.save(docx_path)
        return docx_path
